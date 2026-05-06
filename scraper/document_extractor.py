"""Document text extraction for PDF, Markdown, plain text, DOCX, and EPUB.

All extractors produce markdown-formatted text so the frontend renders
headings, bold, italic, links, lists, and code blocks uniformly.
"""

import os
import re


SUPPORTED_TYPES = {
    '.pdf': 'pdf',
    '.md': 'markdown',
    '.markdown': 'markdown',
    '.txt': 'text',
    '.docx': 'docx',
    '.epub': 'epub',
}


def detect_file_type(file_path):
    """Return file type string from extension, or None if unsupported."""
    ext = os.path.splitext(file_path)[1].lower()
    return SUPPORTED_TYPES.get(ext)


def extract_document(file_path, file_type):
    """Extract text sections from a document.

    Returns list of {index, label, text} dicts.
    label is None for flat/unstructured documents.
    """
    handlers = {
        'pdf': _extract_pdf,
        'markdown': _extract_markdown,
        'text': _extract_text,
        'docx': _extract_docx,
        'epub': _extract_epub,
    }
    handler = handlers.get(file_type)
    if not handler:
        raise ValueError(f"Unsupported file type: {file_type}")
    sections = handler(file_path)
    for i, s in enumerate(sections):
        s['index'] = i
    return sections


def _normalize_pdf_text(text):
    """Join soft-wrapped PDF lines into paragraphs, then apply heading heuristics."""
    text = re.sub(r'\n{3,}', '\n\n', text)
    paragraphs = text.split('\n\n')
    result = []
    for para in paragraphs:
        lines = para.split('\n')
        joined_lines = []
        for line in lines:
            line = line.rstrip()
            if not line:
                continue
            if joined_lines and not re.search(r'[.!?:;"—]\s*$', joined_lines[-1]):
                joined_lines[-1] = joined_lines[-1].rstrip('-') + (
                    '' if joined_lines[-1].endswith('-') else ' '
                ) + line.lstrip()
            else:
                joined_lines.append(line)
        if joined_lines:
            result.append('\n'.join(joined_lines))

    md_parts = []
    for para in result:
        stripped = para.strip()
        if (_is_heading_like(stripped)):
            md_parts.append(f'## {stripped}')
        else:
            md_parts.append(stripped)
    return '\n\n'.join(md_parts)


def _is_heading_like(line):
    """Heuristic: short, title-cased or all-caps lines are probably headings."""
    if not line or len(line) > 80:
        return False
    if line.endswith(('.', ',', ';', ':', '—')):
        return False
    words = line.split()
    if len(words) < 2 or len(words) > 10:
        return False
    if re.match(r'^page\s+\d+$', line, re.IGNORECASE):
        return False
    if line.isupper():
        return True
    if line.istitle() and len(line) < 60:
        return True
    return False


def _extract_pdf(file_path):
    try:
        import fitz
    except ImportError:
        raise ImportError("PyMuPDF required: pip install pymupdf")

    doc = fitz.open(file_path)
    toc = doc.get_toc()
    sections = []

    if toc:
        page_ranges = []
        for i, (_, title, page) in enumerate(toc):
            start = page - 1
            end = toc[i + 1][2] - 2 if i + 1 < len(toc) else len(doc) - 1
            page_ranges.append((title, start, max(start, end)))

        for idx, (title, start, end) in enumerate(page_ranges):
            text = ''.join(doc[pg].get_text() for pg in range(start, min(end + 1, len(doc))))
            text = _normalize_pdf_text(text.strip())
            if text:
                sections.append({'index': idx, 'label': title, 'text': text})
    else:
        # No TOC — extract all pages then split on detected headings
        all_text = ''
        for pg in range(len(doc)):
            all_text += doc[pg].get_text() + '\n\n'
        all_text = _normalize_pdf_text(all_text.strip())

        if all_text:
            heading_pattern = re.compile(r'^## (.+)$', re.MULTILINE)
            parts = heading_pattern.split(all_text)
            if len(parts) > 1:
                if parts[0].strip():
                    sections.append({'index': 0, 'label': None, 'text': parts[0].strip()})
                for i in range(1, len(parts), 2):
                    label = parts[i].strip()
                    body = parts[i + 1].strip() if i + 1 < len(parts) else ''
                    if body:
                        sections.append({'index': 0, 'label': label, 'text': body})
                    elif label:
                        sections.append({'index': 0, 'label': label, 'text': ''})
            else:
                sections.append({'index': 0, 'label': None, 'text': all_text})

    doc.close()
    return sections


def _extract_markdown(file_path):
    """Preserve raw markdown — split on top-level headings for sections."""
    with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
        content = f.read()

    lines = content.split('\n')
    sections = []
    current_label = None
    current_lines = []

    def _flush(label, buf):
        text = '\n'.join(buf).strip()
        if text:
            sections.append({'index': 0, 'label': label, 'text': text})

    for line in lines:
        if re.match(r'^#{1,2}\s+', line):
            _flush(current_label, current_lines)
            current_label = re.sub(r'^#+\s+', '', line).strip()
            current_lines = []
        else:
            current_lines.append(line)

    _flush(current_label, current_lines)

    if not sections:
        sections = [{'index': 0, 'label': None, 'text': content.strip()}]

    return sections


def _extract_text(file_path):
    with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
        content = f.read()
    return [{'index': 0, 'label': None, 'text': content.strip()}]


def _extract_docx(file_path):
    """Extract DOCX with inline formatting preserved as markdown."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx required: pip install python-docx")

    doc = Document(file_path)
    sections = []
    current_label = None
    current_paras = []

    def _para_to_md(para):
        """Convert a python-docx paragraph to markdown with inline formatting."""
        parts = []
        for run in para.runs:
            t = run.text
            if not t:
                continue
            if run.bold and run.italic:
                t = f'***{t}***'
            elif run.bold:
                t = f'**{t}**'
            elif run.italic:
                t = f'*{t}*'
            parts.append(t)
        line = ''.join(parts)
        style = para.style.name
        if style.startswith('List'):
            line = f'- {line}'
        return line

    def _flush(label, paras):
        text = '\n\n'.join(paras).strip()
        if text:
            sections.append({'index': 0, 'label': label, 'text': text})

    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            _flush(current_label, current_paras)
            current_label = para.text.strip()
            current_paras = []
        elif para.text.strip():
            current_paras.append(_para_to_md(para))

    _flush(current_label, current_paras)

    if not sections:
        text = '\n\n'.join(_para_to_md(p) for p in doc.paragraphs if p.text.strip())
        sections = [{'index': 0, 'label': None, 'text': text}]

    return sections


def _extract_epub(file_path):
    """Convert EPUB chapters from HTML to markdown."""
    try:
        from ebooklib import epub, ITEM_DOCUMENT
        from bs4 import BeautifulSoup
        import html2text
    except ImportError:
        raise ImportError("ebooklib + beautifulsoup4 + html2text required")

    h = html2text.HTML2Text()
    h.body_width = 0
    h.ignore_images = True
    h.unicode_snob = True

    book = epub.read_epub(file_path, options={'ignore_ncx': True})
    sections = []

    for item in book.get_items_of_type(ITEM_DOCUMENT):
        content = item.get_content()
        soup = BeautifulSoup(content, 'html.parser')
        for tag in soup.find_all(['script', 'style']):
            tag.decompose()
        md = h.handle(str(soup)).strip()
        if not md:
            continue
        heading = soup.find(['h1', 'h2', 'h3'])
        label = heading.get_text(strip=True) if heading else item.get_name()
        sections.append({'index': 0, 'label': label, 'text': md})

    return sections
